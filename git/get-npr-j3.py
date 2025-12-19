import os
from googletrans import Translator
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import html


def translate_and_export(txt_file: str, mp3_filename: str, dest_folder: str = ".", verbose: bool = True):
    """
    Đọc transcript .txt, dịch sang tiếng Việt bằng Google Translate,
    rồi xuất ra DOCX + PDF song ngữ.
    """
    # --- Đọc transcript ---
    with open(txt_file, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    en_lines = [line.strip() for line in lines if line.strip()]

    # --- Google Translate ---
    translator = Translator()
    vi_lines = []
    for line in en_lines:
        try:
            vi = translator.translate(line, dest="vi").text
        except Exception as e:
            print("❌ Google Translate lỗi:", e)
            raise RuntimeError("Không dịch được, script dừng.")
        vi_lines.append(vi)

    # --- Xuất DOCX ---
    base, _ = os.path.splitext(mp3_filename)
    docx_path = os.path.join(dest_folder, f"{base}.docx")

    doc = Document()
    doc.add_heading("NPR Transcript (EN + VI)", 0)
    for en, vi in zip(en_lines, vi_lines):
        doc.add_paragraph(en)
        doc.add_paragraph(vi, style="Intense Quote")
        doc.add_paragraph("")
    doc.save(docx_path)
    if verbose:
        print("✅ Đã tạo DOCX:", docx_path)

    # --- Xuất PDF ---
    pdf_path = os.path.join(dest_folder, f"{base}.pdf")
    styles = getSampleStyleSheet()
    story = []
    for en, vi in zip(en_lines, vi_lines):
        story.append(Paragraph(html.escape(en), styles["Normal"]))
        story.append(Paragraph(html.escape(vi), styles["Normal"]))
        story.append(Spacer(1, 12))

    pdf = SimpleDocTemplate(pdf_path)
    pdf.build(story)
    if verbose:
        print("✅ Đã tạo PDF:", pdf_path)

    return docx_path, pdf_path


# --- Ví dụ sử dụng ---
if __name__ == "__main__":
    mp3_file = "npr_audio.mp3"   # từ Job 1
    txt_file = "npr_audio.txt"   # từ Job 2
    translate_and_export(txt_file, mp3_file)
